#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民间借贷纠纷答辩状适配器 - 端到端测试

修复v2：正确处理"无□/有□"两段落结构
- para0: '无□' → 选"无"时勾选无■，不填内容
- para1: '有□    异议内容：' → 选"有"时勾选有■，内容填在para1末尾
"""

import os, sys, shutil
from datetime import datetime
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from engine.template_engine import BookmarkEngine
from adapters.civil_lending_defense import CivilLendingDefenseAdapter

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(PROJECT_ROOT, "template_cache")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
TEMPLATE_FILE = "民事答辩状-民间借贷纠纷.docx"
OUTPUT_FILE = f"民事答辩状-民间借贷纠纷测试_李四答辩_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"


def add_defense_bookmarks(engine, case_data):
    """为答辩状模板添加书签
    
    关键：答辩状模板每个异议项的C1单元格有2个段落：
    - para0: '无□'
    - para1: '有□    异议内容：'
    
    书签添加策略：
    - 对于"有异议=True"的项：在para1末尾添加书签（异议内容/事实与理由后面）
    - 对于"有异议=False"的项：不添加内容书签，只通过勾选框处理"无■"
    """
    
    # ============ T0 ============
    # T0 Row1 C0: 案号
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=0, row_index=1, cell_index=0,
        bookmark_prefix="D_案号",
        field_names=["案号"],
        paragraph_indices=[0]
    )
    
    # T0 Row3 C1: 答辩人(自然人) - 8个段落
    # para0: '姓名：'
    # para1: '性别：男□    女□'
    # para2: '出生日期：      年        月         日                            民族：'
    # para3: '工作单位：                        职务：                        联系电话：'
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=0, row_index=3, cell_index=1,
        bookmark_prefix="D_答辩人",
        field_names=["姓名：", "出生日期：", "民族：", "工作单位：", "职务：", "联系电话："],
        paragraph_indices=[0, 2, 2, 3, 3, 3]
    )
    
    # ============ T1 ============
    # T1 Row0 C1: 委托代理人
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=1, row_index=0, cell_index=1,
        bookmark_prefix="D_代理人",
        field_names=["姓名：", "单位："],
        paragraph_indices=[1, 2]
    )
    
    # T1 Row3-Row10: 诉讼请求异议(8项)
    # 关键：只对"有异议=True"的项添加书签
    # 模板结构：para0='无□', para1='有□    异议内容：'
    # 书签需要添加到para1（index=1）末尾
    objection_rows_t1 = [
        (3, "本金", case_data.get("诉讼请求异议", {}).get("本金", {})),
        (4, "利息", case_data.get("诉讼请求异议", {}).get("利息", {})),
        (5, "提前还款", case_data.get("诉讼请求异议", {}).get("提前还款", {})),
        (6, "担保权利", case_data.get("诉讼请求异议", {}).get("担保权利", {})),
        (7, "债权费用", case_data.get("诉讼请求异议", {}).get("债权费用", {})),
        (8, "诉讼费", case_data.get("诉讼请求异议", {}).get("诉讼费", {})),
        (9, "其他请求", case_data.get("诉讼请求异议", {}).get("其他请求", {})),
        (10, "标的总额", case_data.get("诉讼请求异议", {}).get("标的总额", {})),
    ]
    for row_idx, item, data in objection_rows_t1:
        has_objection = data.get("有异议", False)
        content = data.get("内容", "")
        # 只有"有异议=True"且有内容时才添加书签
        if has_objection and content:
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=1, row_index=row_idx, cell_index=1,
                bookmark_prefix=f"D_T1_{item}",
                field_names=["异议内容："],
                paragraph_indices=[1]  # 书签添加到para1（'有□    异议内容：'）
            )
    
    # T1 Row13: 事实与理由异议1 (合同签订)
    fact_objection1 = case_data.get("事实异议", {}).get("合同签订", {})
    if fact_objection1.get("有异议", False) and fact_objection1.get("内容", ""):
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, row_index=13, cell_index=1,
            bookmark_prefix="D_T1_合同签订",
            field_names=["事实与理由："],
            paragraph_indices=[1]  # 书签添加到para1
        )
    
    # ============ T2 ============
    # T2 Row0-Row16: 事实与理由异议2-18
    # 模板结构：para0='无□', para1='有□    事实与理由：'
    # 书签需要添加到para1（index=1）末尾
    objection_rows_t2 = [
        (0, "签订主体", case_data.get("事实异议", {}).get("签订主体", {})),
        (1, "借款金额", case_data.get("事实异议", {}).get("借款金额", {})),
        (2, "借款期限", case_data.get("事实异议", {}).get("借款期限", {})),
        (3, "借款利率", case_data.get("事实异议", {}).get("借款利率", {})),
        (4, "借款提供时间", case_data.get("事实异议", {}).get("借款提供时间", {})),
        (5, "还款方式", case_data.get("事实异议", {}).get("还款方式", {})),
        (6, "还款情况", case_data.get("事实异议", {}).get("还款情况", {})),
        (7, "逾期还款", case_data.get("事实异议", {}).get("逾期还款", {})),
        (8, "物的担保", case_data.get("事实异议", {}).get("物的担保", {})),
        (9, "担保人担保物", case_data.get("事实异议", {}).get("担保人担保物", {})),
        (10, "最高额担保", case_data.get("事实异议", {}).get("最高额担保", {})),
        (11, "抵押质押登记", case_data.get("事实异议", {}).get("抵押质押登记", {})),
        (12, "保证合同", case_data.get("事实异议", {}).get("保证合同", {})),
        (13, "保证方式", case_data.get("事实异议", {}).get("保证方式", {})),
        (14, "其他担保", case_data.get("事实异议", {}).get("其他担保", {})),
        (15, "免责减责", case_data.get("事实异议", {}).get("免责减责", {})),
        (16, "其他说明", case_data.get("事实异议", {}).get("其他说明", {})),
    ]
    for row_idx, item, data in objection_rows_t2:
        has_objection = data.get("有异议", False)
        content = data.get("内容", "")
        label = "内容：" if item == "其他说明" else "事实与理由："
        # 只有"有异议=True"且有内容时才添加书签
        if has_objection and content:
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=2, row_index=row_idx, cell_index=1,
                bookmark_prefix=f"D_T2_{item}",
                field_names=[label],
                paragraph_indices=[1]  # 书签添加到para1（'有□    事实与理由：'）
            )
    
    # T2 Row17 C1: 答辩依据(合同约定/法律规定，2个段落)
    basis = case_data.get("答辩依据", {})
    if basis.get("合同约定", "") or basis.get("法律规定", ""):
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=2, row_index=17, cell_index=1,
            bookmark_prefix="D_答辩依据",
            field_names=["合同约定：", "法律规定："],
            paragraph_indices=[0, 1]
        )
    
    print("  书签添加完成!")


def main():
    print("=" * 80)
    print("民间借贷纠纷答辩状适配器 - 端到端测试 v2")
    print("=" * 80)

    # 1. 加载适配器
    print("\n[1/5] 加载适配器...")
    adapter = CivilLendingDefenseAdapter()
    print(f"  适配器名称: {adapter.name()}")
    print(f"  模板名称: {adapter.get_template_name()}")

    # 2. 准备测试数据
    print("\n[2/5] 准备模拟答辩案件数据...")
    case_data = {
        "案号": "（2024）冀05民初1234号",
        "答辩人": {
            "类型": "自然人",
            "姓名": "李四",
            "性别": "男",
            "出生日期": "1985年6月15日",
            "民族": "汉族",
            "工作单位": "邢台市某商贸公司",
            "职务": "经理",
            "联系电话": "13903190123",
        },
        "代理人": {
            "姓名": "赵律师",
            "单位": "河北振环律师事务所",
            "职务": "主办律师",
            "联系电话": "0319-12345678",
            "类型": "律师",
        },
        "诉讼请求异议": {
            "本金": {"有异议": False, "内容": "确认收到借款本金40万元（而非原告主张的50万元）。"},
            "利息": {"有异议": True, "内容": "原告主张利息12万元有误，实际应付利息约8万元。"},
            "提前还款": {"有异议": False, "内容": ""},
            "担保权利": {"有异议": False, "内容": "双方未签订任何担保合同。"},
            "债权费用": {"有异议": True, "内容": "原告主张的律师费3万元无合同依据，不应由被告承担。"},
            "诉讼费": {"有异议": False, "内容": "由法院依法酌定。"},
            "其他请求": {"有异议": False, "内容": ""},
            "标的总额": {"有异议": True, "内容": "标的总额应按实际借款40万元及应付利息8万元计算，合计48万元。"},
        },
        "事实异议": {
            "合同签订": {"有异议": False, "内容": ""},
            "签订主体": {"有异议": False, "内容": ""},
            "借款金额": {"有异议": True, "内容": "实际收到借款40万元，原告主张的50万元与事实不符。"},
            "借款期限": {"有异议": False, "内容": ""},
            "借款利率": {"有异议": True, "内容": "双方口头约定年利率10%，原告主张的15%与实际不符。"},
            "借款提供时间": {"有异议": False, "内容": ""},
            "还款方式": {"有异议": False, "内容": ""},
            "还款情况": {"有异议": True, "内容": "已通过银行转账还款10万元（2023年6月1日），原告未计入还款记录。"},
            "逾期还款": {"有异议": True, "内容": "部分认可逾期，但系因资金周转困难，属情有可原。"},
            "物的担保": {"有异议": False, "内容": "无物的担保。"},
            "担保人担保物": {"有异议": False, "内容": ""},
            "最高额担保": {"有异议": False, "内容": ""},
            "抵押质押登记": {"有异议": False, "内容": ""},
            "保证合同": {"有异议": True, "内容": "保证合同系在原告欺骗下签署，并非被告真实意思表示。"},
            "保证方式": {"有异议": False, "内容": ""},
            "其他担保": {"有异议": False, "内容": ""},
            "免责减责": {"有异议": True, "内容": "2023年因不可抗力（自然灾害）导致经营困难，属于免责事由。"},
            "其他说明": {"有异议": False, "内容": ""},
        },
        "答辩依据": {
            "合同约定": "《借款合同》第三条约定借款利率以实际到账金额为准。",
            "法律规定": "《民法典》第六百七十五条、《最高人民法院关于审理民间借贷案件适用法律若干问题的规定》(2020修订)第二十五条。",
        },
        "调解意愿": {
            "了解调解": True,
            "先行调解": "是",
        },
    }
    print(f"  答辩人: {case_data['答辩人']['姓名']} ({case_data['答辩人']['性别']})")
    print(f"  案号: {case_data['案号']}")
    异议_count = sum(1 for v in case_data['诉讼请求异议'].values() if v.get('有异议'))
    print(f"  诉讼请求异议: {异议_count}项有异议")

    # 3. 构建填充数据 - 关键修复：只在"有异议=True"时填入内容
    text_map = {
        # T0 案号和答辩人
        "D_案号_案号": case_data["案号"],
        "D_答辩人_姓名：": case_data["答辩人"]["姓名"],
        "D_答辩人_出生日期：": case_data["答辩人"]["出生日期"],
        "D_答辩人_民族：": case_data["答辩人"]["民族"],
        "D_答辩人_工作单位：": case_data["答辩人"]["工作单位"],
        "D_答辩人_职务：": case_data["答辩人"]["职务"],
        "D_答辩人_联系电话：": case_data["答辩人"]["联系电话"],
        # T1 代理人
        "D_代理人_姓名：": case_data["代理人"]["姓名"],
        "D_代理人_单位：": case_data["代理人"]["单位"],
        # T2 答辩依据
        "D_答辩依据_合同约定：": case_data["答辩依据"]["合同约定"],
        "D_答辩依据_法律规定：": case_data["答辩依据"]["法律规定"],
    }
    
    # T1 诉讼请求异议 - 只有"有异议=True"且有内容时才填入
    for item, data in case_data["诉讼请求异议"].items():
        if data.get("有异议", False) and data.get("内容", ""):
            text_map[f"D_T1_{item}_异议内容："] = data["内容"]
    
    # T1 事实异议1
    fact1 = case_data["事实异议"].get("合同签订", {})
    if fact1.get("有异议", False) and fact1.get("内容", ""):
        text_map[f"D_T1_合同签订_事实与理由："] = fact1["内容"]
    
    # T2 事实异议2-18 - 只有"有异议=True"且有内容时才填入
    for item, data in case_data["事实异议"].items():
        if item == "合同签订":
            continue  # 已在T1处理
        label = "内容：" if item == "其他说明" else "事实与理由："
        if data.get("有异议", False) and data.get("内容", ""):
            text_map[f"D_T2_{item}_{label}"] = data["内容"]
    
    # 过滤空字符串
    text_map = {k: v for k, v in text_map.items() if v}
    
    # 勾选框数据
    checkbox_map = {
        # T0 答辩人性别
        "T0_答辩人_性别_男": case_data["答辩人"]["性别"] == "男",
        "T0_答辩人_性别_女": case_data["答辩人"]["性别"] == "女",
        # T1 委托代理人
        "T1_委托代理人_有": True,
        "T1_委托代理人_类型_律师": case_data["代理人"]["类型"] == "律师",
        "T1_委托代理人_类型_亲友": case_data["代理人"]["类型"] == "亲友",
        "T1_委托代理人_类型_工作人员": case_data["代理人"]["类型"] == "工作人员",
        "T1_委托代理人_类型_基层组织推荐": case_data["代理人"]["类型"] == "基层组织推荐",
        "T1_委托代理人_类型_其他": case_data["代理人"]["类型"] == "其他",
    }
    
    # T1 诉讼请求异议 无/有 - 根据"有异议"值决定勾选
    for item, data in case_data["诉讼请求异议"].items():
        has_objection = data.get("有异议", False)
        checkbox_map[f"T1_{item}_无"] = not has_objection
        checkbox_map[f"T1_{item}_有"] = has_objection
    
    # T1 事实异议1
    fact1 = case_data["事实异议"].get("合同签订", {})
    has_objection1 = fact1.get("有异议", False)
    checkbox_map["T1_合同签订_无"] = not has_objection1
    checkbox_map["T1_合同签订_有"] = has_objection1
    
    # T2 事实异议2-18 无/有
    for item, data in case_data["事实异议"].items():
        if item == "合同签订":
            continue
        has_objection = data.get("有异议", False)
        checkbox_map[f"T2_{item}_无"] = not has_objection
        checkbox_map[f"T2_{item}_有"] = has_objection
    
    # T3 调解意愿
    checkbox_map["T3_了解"] = case_data["调解意愿"]["了解调解"]
    checkbox_map["T3_不了解"] = not case_data["调解意愿"]["了解调解"]
    willing = case_data["调解意愿"]["先行调解"]
    checkbox_map["T3_先行调解_是"] = (willing == "是")
    checkbox_map["T3_先行调解_否"] = (willing == "否")
    checkbox_map["T3_先行调解_暂不确定"] = (willing == "暂不确定")
    
    print(f"  文本字段: {len(text_map)}, 勾选字段: {len(checkbox_map)}")

    # 4. 添加书签
    print("\n[3/5] 添加书签...")
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILE)
    engine = BookmarkEngine(template_path)
    add_defense_bookmarks(engine, case_data)

    # 5. 填充书签
    print("\n[4/5] 填充内容...")
    engine.fill(text_map)
    print(f"  填充了 {len(text_map)} 个文本字段")

    # 填充勾选框
    replaced = engine.fill_checkboxes_squares(checkbox_map)
    print(f"  替换了 {replaced} 个勾选框")

    # 6. 保存
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)
    engine.save(output_path)
    engine.cleanup()
    print(f"\n  输出文件: {output_path}")

    # 7. 验证
    print("\n[5/5] 验证输出...")
    from docx import Document
    doc = Document(output_path)

    checks = [
        ("案号填充", "（2024）冀05民初1234号"),
        ("答辩人姓名", "李四"),
        ("答辩人性别(男)", "男"),
        ("代理人姓名", "赵律师"),
        ("利息异议内容(有异议)", "12万元"),
        ("债权费用异议(有异议)", "律师费"),
        ("标的总额异议(有异议)", "48万元"),
        ("借款金额异议(有异议)", "40万元"),
        ("还款情况异议(有异议)", "10万元"),
        ("保证合同异议(有异议)", "欺骗"),
        ("免责减责异议(有异议)", "不可抗力"),
        ("调解意愿", "了解"),
    ]

    all_ok = True
    for name, keyword in checks:
        found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if keyword in cell.text:
                        found = True
                        break
        status = "✓" if found else "✗"
        if not found:
            all_ok = False
        print(f"  {status} {name}: '{keyword}'")

    # 检查勾选框状态
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text

    filled_squares = all_text.count("■")
    print(f"\n  已填充勾选框(■): {filled_squares}")
    print(f"  未填充勾选框(□): {all_text.count('□')}")

    # 验证"无"的项不应该有内容
    print("\n  验证'无异议'项内容为空:")
    no_objection_checks = [
        ("本金(无异议)", "确认收到"),  # 不应该出现
        ("担保权利(无异议)", "担保合同"),  # 不应该出现
        ("诉讼费(无异议)", "酌定"),  # 不应该出现
    ]
    for name, keyword in no_objection_checks:
        found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 检查"无□"后面是否有内容
                    if "无■" in cell.text and keyword in cell.text:
                        found = True
                        break
        status = "✓ 无内容" if not found else "✗ 有内容(错误)"
        if found:
            all_ok = False
        print(f"    {status} {name}")

    print("\n" + "=" * 80)
    if all_ok:
        print("✓ 所有关键内容验证通过！")
    else:
        print("✗ 部分验证失败，请检查。")
    print("=" * 80)
    return all_ok


if __name__ == "__main__":
    ok = main()
    sys.exit(0 if ok else 1)
