#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
建设工程施工合同纠纷适配器端到端测试

测试案例：承包人起诉发包人拖欠工程款

测试流程：
1. 加载模板和适配器
2. 准备测试数据
3. 为模板添加书签（确保书签名称与fill_map键名匹配）
4. 填充书签和勾选框
5. 保存并验证输出
"""

import os
import sys
import shutil
from datetime import datetime

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from adapters.construction_contract import ConstructionContractAdapter
from engine.template_engine import BookmarkEngine

# 项目根目录
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(PROJECT_ROOT, "template_cache")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
TEMPLATE_FILE = "民事起诉状-建设工程施工合同纠纷.docx"


def create_test_case_data():
    """创建测试案件数据"""
    return {
        # === 原告（承包人）信息 ===
        "plaintiff_type": "法人",
        "plaintiff_name": "深圳市建筑工程有限公司",
        "plaintiff_address": "深圳市福田区莲花路100号",
        "plaintiff_registration_address": "深圳市福田区",
        "plaintiff_legal_rep": "张三",
        "plaintiff_legal_rep_position": "董事长",
        "plaintiff_phone": "13800138000",
        "plaintiff_id_number": "91440300MA5DXXXXX",
        "plaintiff_company_type": "有限责任公司",
        "plaintiff_ownership": "民营",
        
        # === 代理人信息 ===
        "has_agent": True,
        "agent_name": "李四",
        "agent_unit": "广东华律师事务所",
        "agent_position": "律师",
        "agent_phone": "13900139000",
        "agent_authorization": "特别授权",
        
        # === 被告（发包人）信息 ===
        "defendant_type": "法人",
        "defendant_name": "东莞市房地产开发有限公司",
        "defendant_address": "东莞市南城区鸿福路108号",
        "defendant_registration_address": "东莞市工商行政管理局",
        "defendant_legal_rep": "王五",
        "defendant_legal_rep_position": "总经理",
        "defendant_phone": "13700137000",
        "defendant_id_number": "91441900MA4UXXXXX",
        "defendant_company_type": "有限责任公司",
        "defendant_ownership": "民营",
        
        # === 第三人信息 ===
        "has_third_party": False,
        
        # === 原告角色 ===
        "plaintiff_role": "承包人或施工人",
        
        # === 诉讼请求（承包人） ===
        # 1. 支付工程款
        "claim_project_payment": 5000000,
        
        # 2. 迟延支付工程款的利息（违约金）
        "claim_interest_until_date": "2024年6月30日",
        "claim_interest_amount": 200000,
        "claim_penalty_amount": 100000,
        "claim_interest_penalty_from_date": "2024年7月1日",
        "claim_interest_penalty_base": 5000000,
        "claim_interest_penalty_rate": "全国银行间同业拆借中心公布的贷款市场报价利率",
        "claim_interest_to_actual": True,
        
        # 3. 建设工程价款优先受偿权
        "claim_priority_right": True,
        "claim_priority_content": "对涉案工程折价款或拍卖款享有建设工程价款优先受偿权",
        
        # 4. 突破合同相对性
        "claim_bypass_liability": False,
        
        # 5. 赔偿损失
        "claim_loss": True,
        "claim_loss_amount": 300000,
        "claim_loss_type": "停窝工损失",
        "claim_loss_detail": "因被告未按约支付工程进度款，导致原告停工窝工",
        "claim_loss_basis": "停窝工期间的人员工资、设备租赁费用等实际支出",
        
        # 6. 退还超付工程款
        "claim_overpayment": False,
        
        # === 共同项 ===
        # 7. 支付超付工程款利息
        "claim_overpayment_interest": False,
        
        # 8. 修复责任
        "claim_repair": False,
        
        # 9. 赔偿损失（发包人）
        "claim_defect_loss": False,
        
        # 10. 合同无效
        "claim_contract_invalid": False,
        
        # 11. 继续履行或解除
        "claim_continue_or_dissolve": "继续履行",
        "claim_continue_deadline": "30",
        "claim_continue_obligations": "付款",
        
        # 12. 债权实现费用
        "claim_creditor_expenses": True,
        "claim_creditor_expenses_detail": "律师费150000元、差旅费20000元",
        
        # 13. 诉讼费用
        "claim_litigation_fee": True,
        
        # 14. 其他请求
        "claim_other": "",
        
        # 15. 标的总额
        "total_amount": 5650000,
        
        # === 约定管辖和诉前保全 ===
        "has_jurisdiction_agreement": False,
        "has_preservation": False,
        "has_appraisal": True,
        "appraisal_matters": "申请对涉案工程质量及造价进行鉴定",
        
        # === 事实与理由 ===
        # 1. 合同签订情况
        "contract_name": "建设工程施工合同",
        "contract_number": "GC-2022-001",
        "contract_sign_date": "2022年3月15日",
        "contract_sign_place": "深圳市",
        "contract_bidding": True,
        
        # 2. 签订主体
        "employer_name": "东莞市房地产开发有限公司",
        "contractor_name": "深圳市建筑工程有限公司",
        "资质出借企业": "",
        
        # 3. 建设工程情况
        "project_name": "XX花园住宅小区项目",
        "project_location": "东莞市南城区",
        "project_scope": "土建、装饰、水电安装等全部施工内容",
        "project_quality_standard": "合格",
        
        # 4. 工程款及支付方式
        "project_price_type": "固定单价",
        "project_fixed_price": 2800,
        "project_total_price": 8000000,
        
        # 5. 工期
        "project_start_date": "2022年5月1日",
        "project_end_date": "2024年4月30日",
        "project_duration": 730,
        
        # 6. 工程质量标准及竣工验收
        "quality_acceptance_standard": "按国家标准《建筑工程施工质量验收统一标准》GB50300-2013执行",
        
        # 7. 违约金/保证金
        "has_contract_penalty": True,
        "contract_penalty_amount": 500000,
        "contract_penalty_article": "第十条",
        "has_contract_deposit": True,
        "contract_deposit_amount": 800000,
        "contract_deposit_article": "第九条",
        "has_delay_penalty": True,
        "delay_penalty_rate": "0.5%/日",
        "delay_penalty_article": "第十一条",
        
        # 8. 工程款支付情况
        "project_total_price": 8000000,
        "project_paid_amount": 3000000,
        "project_owed_amount": -5000000,
        "project_owed_interest": 200000,
        
        # 9. 工程质量情况
        "quality_qualified": True,
        
        # 10. 建设工程交付情况
        "delivery_delayed": False,
        "actual_delivery_date": "2024年4月25日",
        
        # 11. 停窝工情况
        "has_suspension": True,
        "suspension_loss": 300000,
        
        # 12. 优先受偿权主张情况
        "priority_claimed": False,
        
        # 13. 其他需要说明的内容
        "other_matters": "",
        
        # 14. 请求依据
        "contract_basis": "根据《建设工程施工合同》专用条款第10条",
        "legal_basis": "根据《中华人民共和国民法典》第七百八十八条，建设工程合同是承包人进行工程建设，发包人支付价款的合同",
        
        # === 调解意愿 ===
        "understand_mediation": True,
        "understand_mediation_benefits": True,
        "consider_mediation": "是",
    }


def add_bookmarks_to_template(engine: BookmarkEngine, case_data: dict, fill_map: dict):
    """
    为建设工程施工合同纠纷模板添加书签
    
    注意：书签名称必须与fill_map键名完全匹配！
    
    T2表格为3列结构：
    - C0: 诉讼请求序号标题
    - C1: 诉讼请求标签（重复C0）
    - C2: 待填写的具体内容
    """
    
    # === T0: 原告信息 ===
    plaintiff_type = case_data.get('plaintiff_type', '自然人')
    
    if plaintiff_type == '自然人':
        # T0 Row 2 原告自然人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=0, 
            row_index=2, 
            cell_index=1,
            bookmark_prefix="T0_01",
            field_names=["原告姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "经常居住地", "证件类型", "证件号码"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6, 7]
        )
    else:
        # T0 Row 3 原告法人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=0, 
            row_index=3, 
            cell_index=1,
            bookmark_prefix="T0_01",
            field_names=["原告名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码", "公司类型", "所有制性质"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6]
        )
    
    # === T1: 委托代理人 ===
    has_agent = case_data.get('has_agent', False)
    if has_agent:
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_01",
            field_names=["代理人姓名"],
            paragraph_indices=[1]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_02",
            field_names=["代理人单位"],
            paragraph_indices=[2]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_03",
            field_names=["代理人职务"],
            paragraph_indices=[2]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_04",
            field_names=["代理人联系电话"],
            paragraph_indices=[2]
        )
    
    # === T1: 被告信息 ===
    defendant_type = case_data.get('defendant_type', '自然人')
    
    if defendant_type == '自然人':
        # T1 Row 1 被告自然人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=1, 
            cell_index=1,
            bookmark_prefix="T1_05",
            field_names=["被告姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "经常居住地", "证件类型", "证件号码"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6, 7]
        )
    else:
        # T1 Row 2 被告法人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=2, 
            cell_index=1,
            bookmark_prefix="T1_05",
            field_names=["被告名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码", "公司类型", "所有制性质"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6]
        )
    
    # === T1: 第三人信息 ===
    if case_data.get('has_third_party'):
        tp_type = case_data.get('third_party_type', '自然人')
        if tp_type == '自然人':
            # T1 Row 3 第三人自然人
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=1, 
                row_index=3, 
                cell_index=1,
                bookmark_prefix="T1_14",
                field_names=["第三人姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "证件号码"],
                paragraph_indices=[0, 1, 2, 3, 3, 4, 7]
            )
        else:
            # T1 Row 4 第三人法人
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=1, 
                row_index=4, 
                cell_index=1,
                bookmark_prefix="T1_14",
                field_names=["第三人名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码", "公司类型", "所有制性质"],
                paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6]
            )
    
    # === T2: 诉讼请求 (2列结构！) ===
    # T2 Row 3: 1. 支付工程款 - C1列填写金额
    engine.add_bookmarks_to_row(2, 3, "T2_01", ["支付工程款金额"], cell_index=1)
    
    # T2 Row 4: 2. 迟延支付工程款的利息（违约金）- C1列
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=2, 
        row_index=4, 
        cell_index=1,
        bookmark_prefix="T2_02",
        field_names=["利息截止日期", "利息金额", "违约金金额", "起始日期", "基数", "标准"],
        paragraph_indices=[0, 0, 0, 1, 1, 2]
    )
    
    # T2 Row 5: 3. 建设工程价款优先受偿权 - C1列
    engine.add_bookmarks_to_row(2, 5, "T2_08", ["优先受偿权内容"], cell_index=1)
    
    # T2 Row 6: 4. 突破合同相对性 - C1列
    engine.add_bookmarks_to_row(2, 6, "T2_09", ["责任主体名称"], cell_index=1)
    
    # T2 Row 7: 5. 赔偿损失 - C1列
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=2, 
        row_index=7, 
        cell_index=1,
        bookmark_prefix="T2_10",
        field_names=["赔偿金金额", "具体情形", "损失计算依据"],
        paragraph_indices=[0, 3, 4]
    )
    
    # T2 Row 8: 6. 退还超付工程款 - C1列
    engine.add_bookmarks_to_row(2, 8, "T2_13", ["超付金额"], cell_index=1)
    
    # === T3: 诉讼请求 7-15项 ===
    # T3 Row 0: 7. 支付超付工程款利息
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=0, 
        cell_index=1,
        bookmark_prefix="T3_01",
        field_names=["超付利息截止日期", "超付利息金额", "超付利息起始日期", "超付利息基数", "超付利息计算方式"],
        paragraph_indices=[0, 0, 1, 1, 2]
    )
    
    # T3 Row 1: 8. 修复责任
    engine.add_bookmarks_to_row(3, 1, "T3_06", ["修复数额"], cell_index=1)
    
    # T3 Row 2: 9. 赔偿损失（发包人）
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=2, 
        cell_index=1,
        bookmark_prefix="T3_07",
        field_names=["赔偿金金额", "具体情形"],
        paragraph_indices=[0, 3]
    )
    
    # T3 Row 3: 10. 合同无效
    engine.add_bookmarks_to_row(3, 3, "T3_09", ["合同无效理由"], cell_index=1)
    
    # T3 Row 4: 11. 继续履行或解除
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=4, 
        cell_index=1,
        bookmark_prefix="T3_10",
        field_names=["继续履行期限", "继续履行义务", "解除合同日期"],
        paragraph_indices=[0, 0, 1]
    )
    
    # T3 Row 5: 12. 债权实现费用
    engine.add_bookmarks_to_row(3, 5, "T3_13", ["债权费用明细"], cell_index=1)
    
    # T3 Row 7: 14. 其他请求
    engine.add_bookmarks_to_row(3, 7, "T3_14", ["其他请求"], cell_index=1)
    
    # T3 Row 8: 15. 标的总额
    engine.add_bookmarks_to_row(3, 8, "T3_15", ["标的总额"], cell_index=1)
    
    # === T3: 约定管辖和诉前保全 ===
    # T3 Row 10: 1. 有无管辖约定
    engine.add_bookmarks_to_row(3, 10, "T3_16", ["管辖条款内容"], cell_index=1)
    
    # T3 Row 11: 2. 是否诉前保全
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=11, 
        cell_index=1,
        bookmark_prefix="T3_17",
        field_names=["保全法院", "保全时间", "案号"],
        paragraph_indices=[0, 0, 1]
    )
    
    # T3 Row 12: 3. 是否申请鉴定
    engine.add_bookmarks_to_row(3, 12, "T3_20", ["鉴定事项"], cell_index=1)
    
    # === T4: 合同详情 ===
    # T4 Row 0: 1. 合同签订情况
    engine.add_bookmarks_to_row(4, 0, "T4_01", ["合同签订情况"], cell_index=1)
    
    # T4 Row 1: 2. 签订主体
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=1, 
        cell_index=1,
        bookmark_prefix="T4_02",
        field_names=["发包人", "承包人", "出借资质企业"],
        paragraph_indices=[0, 1, 2]
    )
    
    # T4 Row 2: 3. 建设工程情况
    engine.add_bookmarks_to_row(4, 2, "T4_05", ["建设工程情况"], cell_index=1)
    
    # T4 Row 3: 4. 工程款及支付方式
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=3, 
        cell_index=1,
        bookmark_prefix="T4_06",
        field_names=["综合单价", "固定单价", "其他计价方式"],
        paragraph_indices=[0, 0, 2]
    )
    
    # T4 Row 4: 5. 工期
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=4, 
        cell_index=1,
        bookmark_prefix="T4_09",
        field_names=["开工时间", "竣工时间", "工期天数"],
        paragraph_indices=[0, 1, 2]
    )
    
    # T4 Row 5: 6. 质量标准及竣工验收
    engine.add_bookmarks_to_row(4, 5, "T4_12", ["质量验收标准"], cell_index=1)
    
    # T4 Row 6: 7. 违约金/保证金
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=6, 
        cell_index=1,
        bookmark_prefix="T4_13",
        field_names=["违约金金额", "违约金条款", "保证金金额", "保证金条款", "迟延违约金比例", "迟延违约金条款"],
        paragraph_indices=[0, 0, 1, 1, 2, 2]
    )
    
    # T4 Row 7: 8. 工程款支付情况
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=7, 
        cell_index=1,
        bookmark_prefix="T4_19",
        field_names=["工程总价", "已支付工程款", "欠超付工程款", "欠超付工程款利息"],
        paragraph_indices=[0, 0, 1, 2]
    )
    
    # T4 Row 8: 9. 工程质量情况
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=8, 
        cell_index=1,
        bookmark_prefix="T4_23",
        field_names=["工程质量问题", "工程质量造成损失"],
        paragraph_indices=[1, 2]
    )
    
    # T4 Row 9: 10. 交付情况
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=9, 
        cell_index=1,
        bookmark_prefix="T4_25",
        field_names=["实际交付时间", "迟延交付造成损失"],
        paragraph_indices=[0, 1]
    )
    
    # T4 Row 10: 11. 停窝工情况
    engine.add_bookmarks_to_row(4, 10, "T4_27", ["停窝工造成损失"], cell_index=1)
    
    # T4 Row 11: 12. 优先受偿权主张
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=11, 
        cell_index=1,
        bookmark_prefix="T4_28",
        field_names=["优先受偿权主张时间", "优先受偿权主张方式"],
        paragraph_indices=[0, 1]
    )
    
    # T4 Row 12: 13. 其他需要说明的内容
    engine.add_bookmarks_to_row(4, 12, "T4_30", ["其他说明内容"], cell_index=1)
    
    # T4 Row 13: 14. 请求依据
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=13, 
        cell_index=1,
        bookmark_prefix="T4_31",
        field_names=["合同约定依据", "法律规定依据"],
        paragraph_indices=[0, 1]
    )
    
    print("书签添加完成！")
    print(f"所有书签数量: {len(engine.list_bookmarks())}")


def main():
    """主函数"""
    print("\n" + "="*80)
    print("建设工程施工合同纠纷适配器 - 端到端测试")
    print("="*80)
    
    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 加载适配器
    print("\n[1/5] 加载适配器...")
    adapter = ConstructionContractAdapter()
    print(f"  适配器名称: {adapter.name()}")
    print(f"  模板名称: {adapter.get_template_name()}")
    
    # 检查模板文件
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILE)
    if not os.path.exists(template_path):
        print(f"  错误: 模板文件不存在: {template_path}")
        return False
    
    # 复制模板到输出目录
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(OUTPUT_DIR, f"民事起诉状-建设工程施工合同纠纷测试_{timestamp}.docx")
    shutil.copy(template_path, output_file)
    print(f"  输出文件: {output_file}")
    
    # 创建BookmarkEngine
    print("\n[2/5] 初始化引擎...")
    engine = BookmarkEngine(output_file)
    
    # 生成测试数据
    print("\n[3/5] 生成测试数据...")
    case_data = create_test_case_data()
    print(f"  原告: {case_data['plaintiff_name']}")
    print(f"  被告: {case_data['defendant_name']}")
    print(f"  标的额: {case_data['total_amount']}元")
    
    # 预处理案件数据
    case_data = adapter.analyze_case_data(case_data)
    
    # 生成fill_map
    calc_result = adapter.calculate(case_data)
    fill_map = adapter.build_fill_map(case_data, calc_result)
    print(f"  fill_map字段数: {len(fill_map)}")
    
    # 添加书签（使用与fill_map匹配的名称）
    print("\n[4/5] 添加书签...")
    add_bookmarks_to_template(engine, case_data, fill_map)
    
    # 生成checkbox_map
    checkbox_map = adapter.get_checkbox_map(case_data)
    print(f"  checkbox_map字段数: {len(checkbox_map)}")
    
    # 填充书签
    print("  填充书签...")
    engine.fill(fill_map)
    
    # 填充勾选框
    print("  填充勾选框...")
    engine.fill_checkbox(checkbox_map)
    
    # 保存文件
    print("\n[5/5] 保存文件...")
    engine.save(output_file)
    engine.cleanup()
    
    print(f"\n文件已保存: {output_file}")
    
    # 验证输出
    print("\n验证输出...")
    from docx import Document
    doc = Document(output_file)
    
    # 检查关键内容
    verification_results = []
    
    # 1. 检查原告名称
    plaintiff_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "深圳市建筑工程有限公司" in cell.text:
                    plaintiff_found = True
                    break
    verification_results.append(("原告名称", plaintiff_found))
    
    # 2. 检查被告名称
    defendant_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "东莞市房地产开发有限公司" in cell.text:
                    defendant_found = True
                    break
    verification_results.append(("被告名称", defendant_found))
    
    # 3. 检查工程款金额
    amount_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "5000000" in cell.text:
                    amount_found = True
                    break
    verification_results.append(("工程款金额", amount_found))
    
    # 4. 检查合同名称
    contract_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "建设工程施工合同" in cell.text:
                    contract_found = True
                    break
    verification_results.append(("合同名称", contract_found))
    
    # 5. 检查代理人
    agent_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "李四" in cell.text and "律师" in cell.text:
                    agent_found = True
                    break
    verification_results.append(("代理人信息", agent_found))
    
    # 6. 检查勾选框
    checkbox_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '✓' in cell.text or '√' in cell.text or '■' in cell.text:
                    checkbox_found = True
                    break
    verification_results.append(("勾选框替换", checkbox_found))
    
    # 7. 检查标的总额
    total_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "5650000" in cell.text:
                    total_found = True
                    break
    verification_results.append(("标的总额", total_found))
    
    # 8. 检查工程名称
    project_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "XX花园" in cell.text:
                    project_found = True
                    break
    verification_results.append(("工程名称", project_found))
    
    # 打印验证结果
    print("\n验证结果:")
    print("-" * 40)
    all_passed = True
    for name, result in verification_results:
        status = "✓ 通过" if result else "✗ 失败"
        print(f"  {name}: {status}")
        if not result:
            all_passed = False
    
    print("-" * 40)
    if all_passed:
        print("\n✓ 所有验证通过！测试成功。")
    else:
        print("\n✗ 部分验证失败，请检查。")
    
    return all_passed


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
