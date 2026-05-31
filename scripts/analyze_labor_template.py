# -*- coding: utf-8 -*-
"""
分析劳动争议起诉状模板结构

用法：python scripts/analyze_labor_template.py
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx.oxml.ns import qn


def analyze_template():
    template_path = os.path.join(
        os.path.dirname(__file__), '..', 'template_cache',
        '民事起诉状-劳动争议纠纷.docx'
    )
    template_path = os.path.abspath(template_path)
    print(f"分析模板: {template_path}")
    
    doc = Document(template_path)
    
    print(f"\n=== 文档基本信息 ===")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    
    # 分析所有段落
    print(f"\n=== 段落内容 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"段落 {i}: {text[:100]}...")
    
    # 分析所有表格
    for t_idx, table in enumerate(doc.tables):
        print(f"\n=== 表格 {t_idx} ===")
        print(f"行数: {len(table.rows)}")
        print(f"列数: {len(table.columns)}")
        
        for r_idx, row in enumerate(table.rows):
            print(f"\n行 {r_idx}:")
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  单元格 {c_idx}: {text[:200]}...")
                    
                    # 检查是否有勾选框
                    if '□' in text or '☐' in text:
                        print(f"    [包含勾选框]")
    
    # 检查书签
    print(f"\n=== 检查书签 ===")
    # 获取所有书签
    bookmarks = []
    for element in doc.element.iter():
        if element.tag.endswith('}bookmarkStart'):
            name = element.get('name')
            if name:
                bookmarks.append(name)
    print(f"书签数量: {len(bookmarks)}")
    for bm in bookmarks[:20]:
        print(f"  - {bm}")
    if len(bookmarks) > 20:
        print(f"  ... 还有 {len(bookmarks) - 20} 个书签")


if __name__ == '__main__':
    analyze_template()