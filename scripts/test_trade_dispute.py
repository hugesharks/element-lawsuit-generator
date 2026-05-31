#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
买卖合同纠纷适配器端到端测试

测试案例：卖方起诉买方拖欠货款

测试流程：
1. 加载模板和适配器
2. 准备测试数据
3. 为模板添加书签（确保书签名称与fill_map键名匹配）
4. 填充书签和勾选框
5. 保存并验证输出
"""

import os
import sys
import shutil
from datetime import datetime

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from adapters.trade_dispute import TradeDisputeAdapter
from engine.template_engine import BookmarkEngine

# 项目根目录
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(PROJECT_ROOT, "template_cache")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
TEMPLATE_FILE = "民事起诉状-买卖合同纠纷.docx"


def create_test_case_data():
    """创建测试案件数据"""
    return {
        # === 原告（卖方）信息 ===
        "plaintiff_type": "法人",
        "plaintiff_name": "深圳市华盛贸易有限公司",
        "plaintiff_address": "深圳市南山区科技园路100号",
        "plaintiff_registration_address": "深圳市南山区",
        "plaintiff_legal_rep": "张三",
        "plaintiff_legal_rep_position": "总经理",
        "plaintiff_phone": "13800138000",
        "plaintiff_id_number": "91440300MA5DXXXXX",
        "plaintiff_company_type": "有限责任公司",
        "plaintiff_ownership": "民营",
        
        # === 代理人信息 ===
        "has_agent": True,
        "agent_name": "李四",
        "agent_unit": "广东华律师事务所",
        "agent_position": "律师",
        "agent_phone": "13900139000",
        "agent_authorization": "特别授权",
        
        # === 被告（买方）信息 ===
        "defendant_type": "法人",
        "defendant_name": "东莞市明达电子有限公司",
        "defendant_address": "东莞市松山湖高新区工业园路88号",
        "defendant_registration_address": "东莞市工商行政管理局",
        "defendant_legal_rep": "王五",
        "defendant_legal_rep_position": "董事长",
        "defendant_phone": "13700137000",
        "defendant_id_number": "91441900MA4UXXXXX",
        "defendant_company_type": "有限责任公司",
        "defendant_ownership": "民营",
        
        # === 第三人信息 ===
        "has_third_party": False,
        
        # === 原告诉讼角色 ===
        "plaintiff_role": "卖方",  # 原告作为卖方起诉买方
        
        # === 诉讼请求 ===
        "claim_price_amount": 500000,
        "claim_interest_until_date": "2024年6月30日",
        "claim_interest_amount": 25000,
        "claim_penalty_amount": 10000,
        "claim_interest_penalty_from_date": "2024年7月1日",
        "claim_interest_penalty_base": 500000,
        "claim_interest_penalty_rate": "年利率4.35%",
        "claim_interest_penalty_to_actual": True,
        "claim_defect_responsibility": "否",
        "claim_continue_perform": "继续履行",
        "claim_continue_fulfill_type": "付款",
        "claim_guarantee_rights": False,
        "claim_creditor_expenses": True,
        "claim_creditor_expenses_detail": "律师费30000元、差旅费5000元",
        "claim_litigation_fee": True,
        "total_amount": 560000,
        
        # === 约定管辖和诉前保全 ===
        "has_jurisdiction_agreement": True,
        "jurisdiction_clause": "双方约定如发生争议，由合同签订地深圳市南山区人民法院管辖",
        "has_preservation": False,
        
        # === 事实与理由 ===
        "contract_name": "产品购销合同",
        "contract_number": "XS-2024-001",
        "contract_sign_date": "2024年1月15日",
        "contract_sign_place": "深圳市南山区",
        "seller_name": "深圳市华盛贸易有限公司",
        "buyer_name": "东莞市明达电子有限公司",
        "goods_description": "电脑配件一批，包括键盘500个（单价100元/个）、鼠标500个（单价50元/个）、显示器200台（单价1500元/台），共计750000元",
        "unit_price": 0,
        "total_price": 750000,
        "payment_method": "转账",
        "payment_schedule": "一次性",
        "delivery_terms": "交货时间：2024年2月28日前；交货地点：买方仓库；运输方式：买方自提；风险承担：货物交付前由卖方承担，交付后由买方承担",
        "quality_standard": "按国家标准GB/T 2828.1-2012验收，质量异议期限为收货后15日内",
        "contract_penalty_article": "第八条",
        "contract_penalty_amount": 0,
        "contract_deposit_article": "第七条",
        "contract_deposit_amount": 75000,
        
        # === 履行情况 ===
        "payment_delivery_status": "卖方已于2024年2月25日按合同约定交付全部货物，并经买方签收确认。买方仅支付部分货款250000元，尚欠货款500000元经多次催告仍未支付",
        "has_delay": True,
        "delay_time": "180天",
        "delay_type": "逾期付款",
        "has_reminder": True,
        "reminder_date": "2024年4月1日",
        "reminder_method": "书面催告函",
        "has_quality_dispute": False,
        "has_non_conformity": False,
        "has_quality_negotiation": False,
        "has_dissolution_notice": False,
        
        # === 担保信息 ===
        "has_property_guarantee": False,
        "has_maximum_guarantee": False,
        "has_registration": False,
        "has_surety": False,
        "has_other_guarantee": False,
        
        # === 请求依据 ===
        "contract_basis": "根据《产品购销合同》第八条约定",
        "legal_basis": "根据《中华人民共和国民法典》第五百七十九条，当事人一方未支付价款、报酬、租金、利息，或者不履行其他金钱债务的，对方可以请求其支付",
        
        # === 调解意愿 ===
        "understand_mediation": True,
        "understand_mediation_benefits": True,
        "consider_mediation": "是",
    }


def add_bookmarks_to_template(engine: BookmarkEngine, case_data: dict, fill_map: dict):
    """
    为买卖合同纠纷模板添加书签
    
    注意：书签名称必须与fill_map键名完全匹配！
    fill_map键名格式示例：T0_01_原告名称, T0_02_原告住所地 等
    """
    
    # === T0: 原告信息 ===
    plaintiff_type = case_data.get('plaintiff_type', '自然人')
    
    if plaintiff_type == '自然人':
        # T0 Row 2 原告自然人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=0, 
            row_index=2, 
            cell_index=1,
            bookmark_prefix="T0_01",
            field_names=["原告姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "经常居住地", "证件类型", "证件号码"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6, 7]
        )
    else:
        # T0 Row 3 原告法人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=0, 
            row_index=3, 
            cell_index=1,
            bookmark_prefix="T0_01",
            field_names=["原告名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码", "公司类型"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5]
        )
    
    # === T1: 委托代理人 ===
    has_agent = case_data.get('has_agent', False)
    if has_agent:
        # 使用正确的书签前缀匹配fill_map键名
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_01",
            field_names=["代理人姓名"],
            paragraph_indices=[1]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_02",
            field_names=["代理人单位"],
            paragraph_indices=[2]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_03",
            field_names=["代理人职务"],
            paragraph_indices=[2]
        )
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=0, 
            cell_index=1,
            bookmark_prefix="T1_04",
            field_names=["代理人联系电话"],
            paragraph_indices=[2]
        )
    
    # === T1: 被告信息 ===
    defendant_type = case_data.get('defendant_type', '自然人')
    has_third_party = case_data.get('has_third_party', False)
    
    if defendant_type == '自然人':
        # T1 Row 1 被告自然人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=1, 
            cell_index=1,
            bookmark_prefix="T1_05",
            field_names=["被告姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "经常居住地", "证件类型", "证件号码"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5, 6, 7]
        )
    else:
        # T1 Row 2 被告法人
        engine.add_bookmarks_to_cell_paragraphs(
            table_index=1, 
            row_index=2, 
            cell_index=1,
            bookmark_prefix="T1_05",
            field_names=["被告名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码", "公司类型"],
            paragraph_indices=[0, 1, 2, 3, 3, 4, 5]
        )
    
    # === T1: 第三人信息 ===
    if has_third_party:
        tp_type = case_data.get('third_party_type', '自然人')
        if tp_type == '自然人':
            # T1 Row 3 第三人自然人
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=1, 
                row_index=3, 
                cell_index=1,
                bookmark_prefix="T1_14",
                field_names=["第三人姓名", "性别", "出生日期", "工作单位职务", "联系电话", "住所地", "证件号码"],
                paragraph_indices=[0, 1, 2, 3, 3, 4, 7]
            )
        else:
            # T1 Row 4 第三人法人
            engine.add_bookmarks_to_cell_paragraphs(
                table_index=1, 
                row_index=4, 
                cell_index=1,
                bookmark_prefix="T1_14",
                field_names=["第三人名称", "住所地", "注册地", "法定代表人职务", "联系电话", "统一社会信用代码"],
                paragraph_indices=[0, 1, 2, 3, 3, 4]
            )
    
    # === T2: 诉讼请求 ===
    # T2 Row 3 给付价款
    engine.add_bookmarks_to_row(2, 3, "T2_01", ["给付价款"], cell_index=1)
    
    # T2 Row 4 迟延利息违约金 - 多段落
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=2, 
        row_index=4, 
        cell_index=1,
        bookmark_prefix="T2_02",
        field_names=["利息截止日期", "利息金额", "违约金金额", "起始日期", "基数", "标准"],
        paragraph_indices=[0, 0, 0, 2, 2, 3]
    )
    
    # T2 Row 5 违约损失
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=2, 
        row_index=5, 
        cell_index=1,
        bookmark_prefix="T2_08",
        field_names=["赔偿金金额", "违约类型", "具体情形", "计算依据"],
        paragraph_indices=[0, 1, 3, 4]
    )
    
    # T2 Row 6 瑕疵责任
    engine.add_bookmarks_to_row(2, 6, "T2_11", ["瑕疵责任其他方式"], cell_index=1)
    
    # T2 Row 7 继续履行或解除
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=2, 
        row_index=7, 
        cell_index=1,
        bookmark_prefix="T2_12",
        field_names=["期限", "义务类型", "解除日期"],
        paragraph_indices=[0, 0, 1]
    )
    
    # T2 Row 8 担保权利 - 跳过，使用默认值
    
    # T2 Row 9 债权费用
    engine.add_bookmarks_to_row(2, 9, "T2_16", ["债权费用明细"], cell_index=1)
    
    # T2 Row 11 其他请求 - 跳过，使用默认值
    
    # T2 Row 12 标的总额
    engine.add_bookmarks_to_row(2, 12, "T2_18", ["标的总额"], cell_index=1)
    
    # === T3: 约定管辖和诉前保全 ===
    engine.add_bookmarks_to_row(3, 1, "T3_01", ["管辖条款内容"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=2, 
        cell_index=1,
        bookmark_prefix="T3_02",
        field_names=["保全法院", "保全时间", "案号"],
        paragraph_indices=[0, 0, 1]
    )
    
    # === T3: 事实与理由 ===
    engine.add_bookmarks_to_row(3, 5, "T3_05", ["合同名称"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=6, 
        cell_index=1,
        bookmark_prefix="T3_06",
        field_names=["出卖人", "买受人"],
        paragraph_indices=[0, 1]
    )
    engine.add_bookmarks_to_row(3, 7, "T3_07", ["标的物情况"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=8, 
        cell_index=1,
        bookmark_prefix="T3_08",
        field_names=["单价", "总价", "票据类型", "分期方式"],
        paragraph_indices=[0, 0, 1, 2]
    )
    engine.add_bookmarks_to_row(3, 9, "T3_09", ["交货条款"], cell_index=1)
    engine.add_bookmarks_to_row(3, 10, "T3_10", ["质量标准"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=3, 
        row_index=11, 
        cell_index=1,
        bookmark_prefix="T3_11",
        field_names=["违约金条款", "违约金金额", "定金条款", "定金金额", "迟延比例", "迟延条款"],
        paragraph_indices=[0, 0, 1, 1, 2, 2]
    )
    
    # === T4: 履行情况 ===
    engine.add_bookmarks_to_row(4, 0, "T4_01", ["价款支付及交付情况"], cell_index=1)
    # T4 Row 2: 催促履行 - 注意：lxml行索引与python-docx不同
    # python-docx行2对应"是否催促过履行"，但lxml行索引为4
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=4, 
        cell_index=1,
        bookmark_prefix="T4_02",
        field_names=["催促日期", "催促方式"],
        paragraph_indices=[0, 0]
    )
    engine.add_bookmarks_to_row(4, 5, "T4_04", ["质量争议具体情况"], cell_index=1)
    engine.add_bookmarks_to_row(4, 6, "T4_05", ["不符合约定具体情况"], cell_index=1)
    engine.add_bookmarks_to_row(4, 7, "T4_06", ["质量协商具体情况"], cell_index=1)
    engine.add_bookmarks_to_row(4, 8, "T4_07", ["解除合同具体情况"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=9, 
        cell_index=1,
        bookmark_prefix="T4_08",
        field_names=["应付利息", "应付违约金", "应付赔偿金", "共计金额", "计算方式"],
        paragraph_indices=[0, 0, 0, 2, 3]
    )
    engine.add_bookmarks_to_row(4, 10, "T4_13", ["担保合同签订时间"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=11, 
        cell_index=1,
        bookmark_prefix="T4_14",
        field_names=["担保人", "担保物"],
        paragraph_indices=[0, 1]
    )
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=12, 
        cell_index=1,
        bookmark_prefix="T4_16",
        field_names=["担保债权确定时间", "担保额度"],
        paragraph_indices=[1, 2]
    )
    engine.add_bookmarks_to_row(4, 14, "T4_18", ["保证合同签订时间", "保证人", "保证主要内容"], cell_index=1)
    engine.add_bookmarks_to_row(4, 16, "T4_21", ["其他担保形式", "其他担保签订时间"], cell_index=1)
    engine.add_bookmarks_to_cell_paragraphs(
        table_index=4, 
        row_index=17, 
        cell_index=1,
        bookmark_prefix="T4_23",
        field_names=["合同约定依据", "法律规定依据"],
        paragraph_indices=[0, 1]
    )
    
    # === T5: 其他 ===
    engine.add_bookmarks_to_row(5, 0, "T5_01", ["其他需要说明的内容"], cell_index=1)
    engine.add_bookmarks_to_row(5, 1, "T5_02", ["证据清单"], cell_index=1)
    
    print("书签添加完成！")
    print(f"所有书签数量: {len(engine.list_bookmarks())}")


def main():
    """主函数"""
    print("\n" + "="*80)
    print("买卖合同纠纷适配器 - 端到端测试")
    print("="*80)
    
    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 加载适配器
    print("\n[1/5] 加载适配器...")
    adapter = TradeDisputeAdapter()
    print(f"  适配器名称: {adapter.name()}")
    print(f"  模板名称: {adapter.get_template_name()}")
    
    # 检查模板文件
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILE)
    if not os.path.exists(template_path):
        print(f"  错误: 模板文件不存在: {template_path}")
        return False
    
    # 复制模板到输出目录
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(OUTPUT_DIR, f"民事起诉状-买卖合同纠纷测试_{timestamp}.docx")
    shutil.copy(template_path, output_file)
    print(f"  输出文件: {output_file}")
    
    # 创建BookmarkEngine
    print("\n[2/5] 初始化引擎...")
    engine = BookmarkEngine(output_file)
    
    # 生成测试数据
    print("\n[3/5] 生成测试数据...")
    case_data = create_test_case_data()
    print(f"  原告: {case_data['plaintiff_name']}")
    print(f"  被告: {case_data['defendant_name']}")
    print(f"  标的额: {case_data['total_amount']}元")
    
    # 预处理案件数据
    case_data = adapter.analyze_case_data(case_data)
    
    # 生成fill_map
    calc_result = adapter.calculate(case_data)
    fill_map = adapter.build_fill_map(case_data, calc_result)
    print(f"  fill_map字段数: {len(fill_map)}")
    
    # 添加书签（使用与fill_map匹配的名称）
    print("\n[4/5] 添加书签...")
    add_bookmarks_to_template(engine, case_data, fill_map)
    
    # 生成checkbox_map
    checkbox_map = adapter.get_checkbox_map(case_data)
    print(f"  checkbox_map字段数: {len(checkbox_map)}")
    
    # 填充书签
    print("  填充书签...")
    engine.fill(fill_map)
    
    # 填充勾选框
    print("  填充勾选框...")
    engine.fill_checkbox(checkbox_map)
    
    # 保存文件
    print("\n[5/5] 保存文件...")
    engine.save(output_file)
    engine.cleanup()
    
    print(f"\n文件已保存: {output_file}")
    
    # 验证输出
    print("\n验证输出...")
    from docx import Document
    doc = Document(output_file)
    
    # 检查关键内容
    verification_results = []
    
    # 1. 检查原告名称
    plaintiff_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "深圳市华盛贸易有限公司" in cell.text:
                    plaintiff_found = True
                    break
    verification_results.append(("原告名称", plaintiff_found))
    
    # 2. 检查被告名称
    defendant_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "东莞市明达电子有限公司" in cell.text:
                    defendant_found = True
                    break
    verification_results.append(("被告名称", defendant_found))
    
    # 3. 检查标的总额
    amount_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "560000" in cell.text:
                    amount_found = True
                    break
    verification_results.append(("标的总额", amount_found))
    
    # 4. 检查合同名称
    contract_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "产品购销合同" in cell.text:
                    contract_found = True
                    break
    verification_results.append(("合同名称", contract_found))
    
    # 5. 检查代理人
    agent_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "李四" in cell.text and "律师" in cell.text:
                    agent_found = True
                    break
    verification_results.append(("代理人信息", agent_found))
    
    # 6. 检查勾选框
    checkbox_found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '✓' in cell.text or '√' in cell.text or '■' in cell.text:
                    checkbox_found = True
                    break
    verification_results.append(("勾选框替换", checkbox_found))
    
    # 打印验证结果
    print("\n验证结果:")
    print("-" * 40)
    all_passed = True
    for name, result in verification_results:
        status = "✓ 通过" if result else "✗ 失败"
        print(f"  {name}: {status}")
        if not result:
            all_passed = False
    
    print("-" * 40)
    if all_passed:
        print("\n✓ 所有验证通过！测试成功。")
    else:
        print("\n✗ 部分验证失败，请检查。")
    
    return all_passed


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
