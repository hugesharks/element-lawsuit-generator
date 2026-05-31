# -*- coding: utf-8 -*-
"""
赔偿报告生成引擎

核心功能：
- 生成Markdown格式赔偿明细表
- 生成Word格式赔偿报告
- 包含案件信息、赔偿明细、保险拆分、重要提示

使用场景：
- 交通事故案件的赔偿计算报告
- 赔偿项目的详细说明
- 保险理赔的分项计算
"""

from typing import Dict, Any, List, Optional
from datetime import datetime
from dataclasses import dataclass


@dataclass
class CompensationItem:
    """
    赔偿项目数据结构
    
    Attributes:
        name: 项目名称
        amount: 赔偿金额
        formula: 计算公式
        legal_basis: 法律依据
        detail: 详细说明
    """
    name: str
    amount: float
    formula: str = ""
    legal_basis: str = ""
    detail: str = ""


class ReportEngine:
    """
    赔偿/费用报告生成引擎
    
    工作流程：
    1. 接收计算结果数据
    2. 格式化为Markdown表格或Word文档
    3. 包含完整的计算公式和法律依据
    
    使用示例：
        result = {
            "items": [
                {"name": "医疗费", "amount": 52340.80, "formula": "发票金额", "legal_basis": "..."},
            ],
            "total": 123456.78,
            "insurance_split": {...}
        }
        
        md_table = ReportEngine.generate_compensation_table(result)
        ReportEngine.generate_docx_report(result, case_info, "报告.docx")
    """
    
    @staticmethod
    def generate_compensation_table(result: Dict[str, Any], 
                                    standard: Optional[Dict[str, Any]] = None) -> str:
        """
        生成Markdown格式赔偿明细表
        
        Args:
            result: 计算结果字典，包含：
                - items: 赔偿项目列表，每项含name/amount/formula/legal_basis
                - total: 赔偿总额
                - insurance_split: 保险拆分（可选）
            standard: 赔偿标准数据（可选，用于显示标准来源）
            
        Returns:
            Markdown格式的赔偿明细表字符串
        """
        items = result.get("items", [])
        total = result.get("total", 0)
        insurance_split = result.get("insurance_split", {})
        
        lines = [
            "# 交通事故赔偿计算明细",
            "",
            f"**计算日期**: {datetime.now().strftime('%Y年%m月%d日')}",
            "",
        ]
        
        if standard:
            lines.append(f"**适用标准**: 河北省{standard.get('year', '2025')}年度交通事故赔偿标准")
            lines.append("")
        
        lines.extend([
            "## 一、赔偿项目明细",
            "",
            "| 序号 | 赔偿项目 | 金额（元） | 计算公式/依据 |",
            "|:----:|:--------|----------:|:-------------|",
        ])
        
        for i, item in enumerate(items, 1):
            name = item.get("name", "")
            amount = item.get("amount", 0)
            formula = item.get("formula", "")
            # 截断过长的公式
            if len(formula) > 50:
                formula = formula[:47] + "..."
            lines.append(f"| {i} | {name} | {amount:,.2f} | {formula} |")
        
        lines.extend([
            "| | **合计** | **{:,.2f}** | |".format(total),
            "",
        ])
        
        # 保险拆分
        if insurance_split:
            lines.extend([
                "## 二、保险承担分析",
                "",
                "| 保险类型 | 限额（元） | 承担金额（元） | 备注 |",
                "|:--------|----------:|----------:|:-----|",
            ])
            
            compulsory = insurance_split.get("compulsory", {})
            if compulsory:
                lines.append(
                    "| 交强险（死亡伤残） | {:,.2f} | {:,.2f} | |".format(
                        compulsory.get("death_disability_limit", 0),
                        compulsory.get("death_disability_amount", 0)
                    )
                )
                lines.append(
                    "| 交强险（医疗费用） | {:,.2f} | {:,.2f} | |".format(
                        compulsory.get("medical_limit", 0),
                        compulsory.get("medical_amount", 0)
                    )
                )
                lines.append(
                    "| 交强险（财产损失） | {:,.2f} | {:,.2f} | |".format(
                        compulsory.get("property_limit", 0),
                        compulsory.get("property_amount", 0)
                    )
                )
            
            commercial = insurance_split.get("commercial", {})
            if commercial:
                lines.append(
                    "| 商业三者险 | {:,.2f} | {:,.2f} | |".format(
                        commercial.get("limit", 0),
                        commercial.get("amount", 0)
                    )
                )
            
            self_pay = insurance_split.get("self_pay", 0)
            lines.append(
                "| **责任人自付** | - | **{:,.2f}** | |".format(self_pay)
            )
            lines.append("")
        
        # 重要提示
        lines.extend([
            "## 三、重要提示",
            "",
            "1. 以上计算仅供参考，实际赔偿金额以法院判决为准。",
            "2. 赔偿标准依据河北省上年度统计数据，具体以官方公布为准。",
            "3. 医疗费以实际发生且有票据支持的金额为准。",
            "4. 误工费需提供收入证明，否则按当地最低工资标准计算。",
            "5. 伤残等级需经司法鉴定机构鉴定确认。",
            "6. 精神损害抚慰金由法院根据实际情况酌定。",
            "",
            "---",
            "*本报告由要素式法律文书引擎自动生成*",
        ])
        
        return "\n".join(lines)
    
    @staticmethod
    def generate_docx_report(result: Dict[str, Any], 
                             case_info: Dict[str, str],
                             output_path: str,
                             standard: Optional[Dict[str, Any]] = None) -> str:
        """
        生成Word格式赔偿报告
        
        包含：案件信息、赔偿明细（含公式+法律依据）、保险拆分、重要提示
        
        Args:
            result: 计算结果字典
            case_info: 案件信息字典，示例：
                {
                    "案号": "(2024)冀01民初123号",
                    "原告": "张三",
                    "被告": "李四",
                    "事故时间": "2024年1月1日",
                    "责任划分": "主要责任"
                }
            output_path: 输出文件路径
            standard: 赔偿标准数据（可选）
            
        Returns:
            输出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        items = result.get("items", [])
        total = result.get("total", 0)
        insurance_split = result.get("insurance_split", {})
        
        doc = Document()
        
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("交通事故赔偿计算报告")
        run.bold = True
        run.font.size = Pt(18)
        
        # 案件信息
        doc.add_paragraph()
        section_title = doc.add_paragraph()
        run = section_title.add_run("一、案件基本信息")
        run.bold = True
        run.font.size = Pt(14)
        
        if case_info:
            info_table = doc.add_table(rows=len(case_info), cols=2)
            info_table.style = 'Table Grid'
            
            for i, (key, value) in enumerate(case_info.items()):
                info_table.rows[i].cells[0].text = key
                info_table.rows[i].cells[1].text = str(value)
                
                # 设置格式
                for cell in info_table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                
                # 标题列加粗
                for run in info_table.rows[i].cells[0].paragraphs[0].runs:
                    run.bold = True
            
            # 设置列宽
            for row in info_table.rows:
                row.cells[0].width = Cm(3)
                row.cells[1].width = Cm(12)
        
        # 赔偿明细
        doc.add_paragraph()
        section_title = doc.add_paragraph()
        run = section_title.add_run("二、赔偿项目明细")
        run.bold = True
        run.font.size = Pt(14)
        
        # 赔偿明细表格
        detail_table = doc.add_table(rows=1, cols=4)
        detail_table.style = 'Table Grid'
        
        # 表头
        headers = ["序号", "赔偿项目", "金额（元）", "计算公式/法律依据"]
        for i, header in enumerate(headers):
            cell = detail_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        # 数据行
        for idx, item in enumerate(items, 1):
            row = detail_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.get("name", "")
            row.cells[2].text = "{:,.2f}".format(item.get("amount", 0))
            
            # 公式和法律依据合并显示
            formula = item.get("formula", "")
            legal_basis = item.get("legal_basis", "")
            cell_text = formula
            if legal_basis:
                cell_text += f"\n法律依据：{legal_basis}"
            row.cells[3].text = cell_text
            
            # 设置格式
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        # 合计行
        total_row = detail_table.add_row()
        total_row.cells[0].text = ""
        total_row.cells[1].text = "合计"
        total_row.cells[2].text = "{:,.2f}".format(total)
        total_row.cells[3].text = ""
        
        # 合计行加粗
        for cell in total_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        # 设置列宽
        for row in detail_table.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(3)
            row.cells[2].width = Cm(3)
            row.cells[3].width = Cm(8)
        
        # 保险承担分析
        if insurance_split:
            doc.add_paragraph()
            section_title = doc.add_paragraph()
            run = section_title.add_run("三、保险承担分析")
            run.bold = True
            run.font.size = Pt(14)
            
            insurance_table = doc.add_table(rows=1, cols=4)
            insurance_table.style = 'Table Grid'
            
            # 表头
            headers = ["保险类型", "限额（元）", "承担金额（元）", "备注"]
            for i, header in enumerate(headers):
                cell = insurance_table.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            
            # 交强险
            compulsory = insurance_split.get("compulsory", {})
            if compulsory:
                # 死亡伤残
                row = insurance_table.add_row()
                row.cells[0].text = "交强险（死亡伤残）"
                row.cells[1].text = "{:,.2f}".format(compulsory.get("death_disability_limit", 0))
                row.cells[2].text = "{:,.2f}".format(compulsory.get("death_disability_amount", 0))
                row.cells[3].text = ""
                
                # 医疗费用
                row = insurance_table.add_row()
                row.cells[0].text = "交强险（医疗费用）"
                row.cells[1].text = "{:,.2f}".format(compulsory.get("medical_limit", 0))
                row.cells[2].text = "{:,.2f}".format(compulsory.get("medical_amount", 0))
                row.cells[3].text = ""
                
                # 财产损失
                row = insurance_table.add_row()
                row.cells[0].text = "交强险（财产损失）"
                row.cells[1].text = "{:,.2f}".format(compulsory.get("property_limit", 0))
                row.cells[2].text = "{:,.2f}".format(compulsory.get("property_amount", 0))
                row.cells[3].text = ""
            
            # 商业三者险
            commercial = insurance_split.get("commercial", {})
            if commercial:
                row = insurance_table.add_row()
                row.cells[0].text = "商业三者险"
                row.cells[1].text = "{:,.2f}".format(commercial.get("limit", 0))
                row.cells[2].text = "{:,.2f}".format(commercial.get("amount", 0))
                row.cells[3].text = ""
            
            # 责任人自付
            self_pay = insurance_split.get("self_pay", 0)
            row = insurance_table.add_row()
            row.cells[0].text = "责任人自付"
            row.cells[1].text = "-"
            row.cells[2].text = "{:,.2f}".format(self_pay)
            row.cells[3].text = ""
            
            # 自付行加粗
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # 设置格式
            for row in insurance_table.rows[1:]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        
        # 重要提示
        doc.add_paragraph()
        section_title = doc.add_paragraph()
        run = section_title.add_run("四、重要提示")
        run.bold = True
        run.font.size = Pt(14)
        
        tips = [
            "1. 以上计算仅供参考，实际赔偿金额以法院判决为准。",
            "2. 赔偿标准依据河北省上年度统计数据，具体以官方公布为准。",
            "3. 医疗费以实际发生且有票据支持的金额为准。",
            "4. 误工费需提供收入证明，否则按当地最低工资标准计算。",
            "5. 伤残等级需经司法鉴定机构鉴定确认。",
            "6. 精神损害抚慰金由法院根据实际情况酌定。",
        ]
        
        for tip in tips:
            para = doc.add_paragraph(tip)
            for run in para.runs:
                run.font.size = Pt(11)
        
        # 页脚信息
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        run = footer.add_run("\n本报告由要素式法律文书引擎自动生成")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def format_amount(amount: float) -> str:
        """
        格式化金额显示
        
        Args:
            amount: 金额
            
        Returns:
            格式化后的字符串，如 "12,345.67 元"
        """
        return "{:,.2f} 元".format(amount)
    
    @staticmethod
    def generate_summary(result: Dict[str, Any]) -> str:
        """
        生成赔偿计算摘要（纯文本）
        
        Args:
            result: 计算结果字典
            
        Returns:
            摘要文本
        """
        items = result.get("items", [])
        total = result.get("total", 0)
        
        lines = [
            "【赔偿计算摘要】",
            "=" * 40,
        ]
        
        for item in items:
            name = item.get("name", "")
            amount = item.get("amount", 0)
            lines.append(f"• {name}：{amount:,.2f} 元")
        
        lines.extend([
            "=" * 40,
            f"赔偿总额：{total:,.2f} 元",
        ])
        
        insurance_split = result.get("insurance_split", {})
        if insurance_split:
            compulsory = insurance_split.get("compulsory", {})
            compulsory_total = sum([
                compulsory.get("death_disability_amount", 0),
                compulsory.get("medical_amount", 0),
                compulsory.get("property_amount", 0),
            ])
            
            commercial = insurance_split.get("commercial", {})
            commercial_amount = commercial.get("amount", 0)
            
            self_pay = insurance_split.get("self_pay", 0)
            
            lines.extend([
                "-" * 40,
                f"交强险承担：{compulsory_total:,.2f} 元",
                f"商业险承担：{commercial_amount:,.2f} 元",
                f"责任人自付：{self_pay:,.2f} 元",
            ])
        
        return "\n".join(lines)
