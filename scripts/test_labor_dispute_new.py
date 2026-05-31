# -*- coding: utf-8 -*-
"""
劳动争议适配器端到端测试

测试流程：
1. 创建测试案件数据
2. LaborDisputeAdapter 计算赔偿 + 生成映射表
3. BookmarkEngine 加载空白模板
4. 给模板打书签
5. 填充书签内容 + 勾选框
6. 保存输出.docx

运行：cd element-lawsuit-generator && python3 scripts/test_labor_dispute_new.py
"""

import sys
import os
import json

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from adapters.labor_dispute import LaborDisputeAdapter
from engine.template_engine import BookmarkEngine


# ============================================================
# 1. 测试案件数据
# ============================================================
TEST_CASE = {
    # === 原告（劳动者）信息 ===
    "plaintiff_name": "张三",
    "plaintiff_gender": "男",
    "plaintiff_birthdate": "1990年5月15日",
    "plaintiff_ethnicity": "汉族",
    "plaintiff_id_number": "130502199005151234",
    "plaintiff_address": "河北省邢台市襄都区新华南路123号",
    "plaintiff_residence": "河北省邢台市襄都区新华南路123号",
    "plaintiff_work_unit": "邢台某科技有限公司",
    "plaintiff_position": "软件工程师",
    "plaintiff_phone": "13800138001",
    
    # === 被告（用人单位）信息 ===
    "defendant_name": "邢台某科技有限公司",
    "defendant_address": "河北省邢台市襄都区中兴东大街188号",
    "defendant_registration_address": "河北省邢台市襄都区中兴东大街188号",
    "defendant_legal_rep": "李四",
    "defendant_position": "总经理",
    "defendant_phone": "0319-1234567",
    "defendant_credit_code": "91130500MA0G5XKR8T",
    "defendant_company_type": "有限责任公司",
    "defendant_ownership": "民营",
    
    # === 委托代理人 ===
    "has_agent": True,
    "agent_name": "王律师",
    "agent_unit": "河北某律师事务所",
    "agent_position": "律师",
    "agent_phone": "13900139001",
    "agent_authorization": "一般授权",
    
    # === 送达地址 ===
    "delivery_address": "河北省邢台市襄都区法院路1号",
    "delivery_recipient": "张三",
    "delivery_phone": "13800138001",
    "accept_electronic_service": True,
    "electronic_service_method": "短信、微信",
    
    # === 劳动关系基本信息 ===
    "employment_start_date": "2020年3月1日",
    "employment_end_date": "2025年2月28日",
    "work_position": "软件工程师",
    "work_location": "河北省邢台市襄都区",
    "contract_signed": True,
    "contract_sign_date": "2020年3月1日",
    "contract_end_date": "2023年2月28日",
    "contract_type": "固定期限",
    "work_system": "标准工时",
    "monthly_wage": 8000.0,
    "wage_composition": "基本工资6000+绩效2000",
    "actual_monthly_wage": 8000.0,
    "social_insurance_start": "2020年4月",
    "social_insurance_types": "五险",
    
    # === 诉讼请求 ===
    "claim_wage": True,
    "claim_wage_detail": "2025年1月至2月工资16000元",
    "claim_double_wage": False,
    "claim_double_wage_detail": "",
    "claim_overtime": True,
    "claim_overtime_detail": "2024年全年加班费25000元",
    "claim_annual_leave": True,
    "claim_annual_leave_detail": "2023-2024年未休年休假工资7356元",
    "claim_social_insurance_loss": False,
    "claim_social_insurance_loss_detail": "",
    "claim_economic_compensation": True,
    "claim_economic_compensation_detail": "5个月工资40000元",
    "claim_compensation": False,
    "claim_compensation_detail": "",
    "claim_litigation_fee": True,
    "other_claims": "",
    "total_amount": 88356.0,
    
    # === 诉前保全 ===
    "has_preservation": False,
    "preservation_court": "",
    "preservation_time": "",
    "preservation_case_number": "",
    
    # === 事实与理由 ===
    "contract_signing_info": "2020年3月1日，原告与被告签订书面劳动合同，合同期限为2020年3月1日至2023年2月28日，约定月工资8000元。",
    "contract_performance_info": "原告于2020年3月1日入职被告公司，担任软件工程师岗位，工作地点为河北省邢台市襄都区。合同约定月工资8000元（基本工资6000元+绩效2000元），实际按月发放。被告自2020年4月起为原告缴纳五险。",
    "termination_info": "2025年2月28日，被告以公司经营困难为由，单方面解除与原告的劳动合同。原告认为被告解除劳动合同缺乏合法依据，属于违法解除。",
    "work_injury_info": "",
    "arbitration_info": "2025年3月15日，原告向邢台市劳动人事争议仲裁委员会申请仲裁。2025年5月10日，仲裁委作出裁决，但原告对裁决结果不服，依法向人民法院提起诉讼。",
    "other_info": "",
    "legal_basis": "《中华人民共和国劳动合同法》第30条、第46条、第47条、第85条；《中华人民共和国劳动法》第44条；《职工带薪年休假条例》第5条。",
    
    # === 调解意愿 ===
    "understand_mediation": True,
    "understand_mediation_benefits": True,
    "consider_mediation": "是",
    
    # === 工伤相关（本案例无工伤） ===
    "work_injury_date": "",
    "work_injury_recognition": "",
    "disability_level": "",
    "work_injury_medical_expenses": 0.0,
    "hospitalization_start": "",
    "hospitalization_end": "",
}


def main():
    print("=" * 60)
    print("劳动争议适配器端到端测试")
    print("=" * 60)
    
    # --- Step 1: 创建适配器 ---
    print("\n[Step 1] 创建适配器...")
    adapter = LaborDisputeAdapter()
    print(f"  适配器名称: {adapter.name()}")
    
    # --- Step 2: 生成映射表 ---
    print("\n[Step 2] 生成映射表...")
    fill_map = adapter.build_fill_map(TEST_CASE, {})
    print(f"  书签映射数量: {len(fill_map)}")
    
    # 显示书签映射
    print("  书签映射内容:")
    for key, value in list(fill_map.items())[:10]:
        print(f"    {key}: {value[:30]}..." if len(value) > 30 else f"    {key}: {value}")
    if len(fill_map) > 10:
        print(f"    ... 还有 {len(fill_map) - 10} 个映射")
    
    # --- Step 3: 生成勾选框映射 ---
    print("\n[Step 3] 生成勾选框映射...")
    checkbox_map = adapter.get_checkbox_map(TEST_CASE)
    print(f"  勾选框映射数量: {len(checkbox_map)}")
    
    # 显示勾选结果
    checked_count = sum(1 for v in checkbox_map.values() if v)
    print(f"  应勾选数量: {checked_count}")
    
    # --- Step 4: 加载空白模板 ---
    print("\n[Step 4] 加载空白模板...")
    template_path = os.path.join(
        os.path.dirname(__file__), '..', 'template_cache',
        '民事起诉状-劳动争议纠纷.docx'
    )
    template_path = os.path.abspath(template_path)
    print(f"  模板路径: {template_path}")
    
    engine = BookmarkEngine(template_path)
    
    # 列出模板中已有书签
    existing_bookmarks = engine.list_bookmarks()
    print(f"  模板已有书签: {len(existing_bookmarks)}")
    
    # --- Step 5: 插入书签 ---
    print("\n[Step 5] 插入书签...")
    
    # 表格0：当事人信息
    # T0 行2：原告自然人信息（单元格1包含8个段落）
    # 段落0: 姓名 | 段落1: 性别(勾选) | 段落2: 出生日期+民族(多字段) | 段落3: 工作单位+职务+电话(多字段)
    # 段落4: 住所地 | 段落5: 经常居住地 | 段落6: 证件类型 | 段落7: 证件号码
    print("  T0 行2：原告自然人信息...")
    engine.add_bookmarks_to_cell_paragraphs(
        0, 2, 1, "T0",
        ["01_原告姓名", "02_原告住所地", "03_原告经常居住地", "04_原告证件类型", "05_原告证件号码"],
        paragraph_indices=[0, 4, 5, 6, 7]
    )
    
    # T0 行3：委托诉讼代理人（单元格1包含4个段落）
    # 段落0: 有□(勾选) | 段落1: 姓名 | 段落2: 单位+职务+电话(多字段) | 段落3: 代理权限(勾选)
    print("  T0 行3：委托诉讼代理人...")
    engine.add_bookmarks_to_cell_paragraphs(
        0, 3, 1, "T0",
        ["12_代理人姓名", "13_代理人单位", "14_代理人职务", "15_代理人联系电话"],
        paragraph_indices=[1, 2, 2, 2]  # 姓名→para1, 单位/职务/电话→para2
    )
    
    # 表格1：被告信息 + 诉讼请求
    # T1 行0：被告法人信息（Cell 1 包含10个段落）
    # 段落0: 名称 | 段落1: 住所地 | 段落2: 注册地 | 段落3: 法定代表人+职务+电话(多字段)
    # 段落4: 统一社会信用代码 | 段落5-9: 类型勾选
    print("  T1 行0：被告法人信息...")
    engine.add_bookmarks_to_cell_paragraphs(
        1, 0, 1, "T1",
        ["01_被告名称", "02_被告住所地", "03_被告注册地",
         "04_被告法定代表人", "05_被告法定代表人职务", "06_被告联系电话",
         "07_被告统一社会信用代码"],
        paragraph_indices=[0, 1, 2, 3, 3, 3, 4]  # 名称/住所地/注册地/信用代码各一段，法定代表人/职务/电话在para3
    )
    
    # T1 行3-10：8个诉讼请求勾选项（书签名与适配器fill_map一致）
    print("  T1 行3-10：诉讼请求明细...")
    claim_details = [
        (3, "T1_08", "工资支付明细"),
        (4, "T1_09", "双倍工资明细"),
        (5, "T1_10", "加班费明细"),
        (6, "T1_11", "未休年休假工资明细"),
        (7, "T1_12", "社保损失明细"),
        (8, "T1_13", "经济补偿明细"),
        (9, "T1_14", "赔偿金明细"),
        (10, "T1_诉讼费用", ""),  # Row 10 无明细字段，仅勾选框
    ]
    for row_idx, prefix, field in claim_details:
        if field:
            engine.add_bookmarks_to_row(1, row_idx, prefix, [field], cell_index=1)
    
    # T1 行11：其他诉讼请求（空单元格，需要替换内容）
    print("  T1 行11：其他诉讼请求...")
    engine.add_bookmarks_to_row(1, 11, "T1_15", ["其他诉讼请求"], cell_index=1)
    
    # T1 行12：标的总额（空单元格）
    print("  T1 行12：标的总额...")
    engine.add_bookmarks_to_row(1, 12, "T1_16", ["标的总额"], cell_index=1)
    
    # 表格2：事实与理由
    # T2 行4-10：7个事实理由项（书签名与适配器fill_map一致）
    print("  T2 行4-10：事实与理由...")
    fact_items = [
        (4, "T2_04", "劳动合同签订情况"),
        (5, "T2_05", "劳动合同履行情况"),
        (6, "T2_06", "解除终止劳动关系情况"),
        (7, "T2_07", "工伤情况"),
        (8, "T2_08", "劳动仲裁情况"),
        (9, "T2_09", "其他相关情况"),
        (10, "T2_10", "诉请依据"),
    ]
    for row_idx, prefix, field in fact_items:
        engine.add_bookmarks_to_row(2, row_idx, prefix, [field], cell_index=1)
    
    # --- Step 6: 填充书签内容 ---
    print("\n[Step 6] 填充书签内容...")
    filled_count = engine.fill(fill_map)
    print(f"  已填充 {len(fill_map)} 个书签")
    
    # --- Step 7: 填充勾选框 ---
    print("\n[Step 7] 填充勾选框...")
    replaced = engine.fill_checkboxes_squares(checkbox_map)
    print(f"  实际替换 {replaced} 个■")
    
    # --- Step 8: 保存输出文件 ---
    print("\n[Step 8] 保存输出文件...")
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'output')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '测试输出_劳动争议起诉状.docx')
    
    engine.save(output_path)
    
    file_size = os.path.getsize(output_path)
    print(f"  输出路径: {output_path}")
    print(f"  文件大小: {file_size:,} 字节")
    
    # --- Step 9: 验证输出文件 ---
    print("\n[Step 9] 验证输出文件...")
    
    from docx import Document
    doc = Document(output_path)
    
    # 检查关键内容
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += '\n' + cell.text
    
    checks = [
        ("原告姓名", "张三"),
        ("被告名称", "邢台某科技有限公司"),
        ("工资支付明细", "2025年1月至2月工资16000元"),
        ("标的总额", "88356"),
    ]
    
    all_pass = True
    for name, value in checks:
        if value in full_text:
            print(f"  ✅ {name}: 已出现（{value}）")
        else:
            print(f"  ❌ {name}: 未找到（{value}）")
            all_pass = False
    
    # 统计勾选框
    square_count = full_text.count('□')
    black_count = full_text.count('■')
    print(f"\n  📊 勾选框统计: 已勾选■={black_count}, 未勾选□={square_count}")
    
    print("\n" + "=" * 60)
    if all_pass:
        print("✅ 端到端测试完成！")
    else:
        print("⚠️ 部分验证未通过，请检查输出文件。")
    print("=" * 60)
    print(f"\n输出文件: {output_path}")


if __name__ == '__main__':
    main()
