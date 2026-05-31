# -*- coding: utf-8 -*-
"""
证据清单生成引擎

核心功能：
- 定义证据规则（EvidenceRule）
- 注册和管理证据规则
- 根据案件数据自动生成证据清单
- 输出Word或文本格式的证据清单

使用场景：
- 交通事故案件的证据准备指导
- 自动生成证据目录
- 证据完整性检查
"""

from dataclasses import dataclass, field
from typing import Callable, List, Dict, Any, Optional
from datetime import datetime


@dataclass
class EvidenceRule:
    """
    单条证据规则
    
    Attributes:
        name: 证据名称
        category: 类别 - "必备"（必须提供）或"建议补充"（建议提供）
        condition: 条件函数，接收案件数据字典，返回bool
        description: 证明目的描述
        evidence_type: 证据形式（书证/物证/视听资料/电子数据/证人证言/鉴定意见/勘验笔录）
        copies: 份数（默认1份）
        pages: 页数（默认1页）
    """
    name: str
    category: str  # "必备" 或 "建议补充"
    condition: Callable[[Dict[str, Any]], bool]
    description: str
    evidence_type: str = "书证"
    copies: int = 1
    pages: int = 1
    
    def evaluate(self, case_data: Dict[str, Any]) -> bool:
        """
        评估该证据是否需要
        
        Args:
            case_data: 案件数据字典
            
        Returns:
            是否需要该证据
        """
        try:
            return self.condition(case_data)
        except Exception:
            # 条件函数出错时，建议补充类默认不显示
            return self.category == "必备"


class EvidenceEngine:
    """
    证据清单生成引擎
    
    工作流程：
    1. 注册证据规则
    2. 根据案件数据评估每条规则
    3. 生成证据清单
    4. 输出为Word或文本格式
    
    使用示例：
        engine = EvidenceEngine()
        engine.register(EvidenceRule(
            name="身份证",
            category="必备",
            condition=lambda d: True,
            description="证明原告身份信息"
        ))
        
        case_data = {"plaintiff_name": "张三"}
        evidence_list = engine.generate(case_data)
        engine.to_text(evidence_list)
    """
    
    def __init__(self):
        """初始化证据引擎"""
        self.rules: List[EvidenceRule] = []
    
    def register(self, rule: EvidenceRule) -> None:
        """
        注册一条证据规则
        
        Args:
            rule: 证据规则对象
        """
        self.rules.append(rule)
    
    def register_batch(self, rules: List[EvidenceRule]) -> None:
        """
        批量注册证据规则
        
        Args:
            rules: 证据规则列表
        """
        self.rules.extend(rules)
    
    def clear(self) -> None:
        """清空所有已注册的规则"""
        self.rules.clear()
    
    def generate(self, case_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        根据案件数据生成证据清单
        
        遍历所有规则，调用condition判断是否需要该证据
        
        Args:
            case_data: 案件数据字典
            
        Returns:
            证据清单列表，每项包含：
            - name: 证据名称
            - category: 类别（必备/建议补充）
            - description: 证明目的
            - evidence_type: 证据形式
            - copies: 份数
            - pages: 页数
        """
        evidence_list = []
        
        for rule in self.rules:
            if rule.evaluate(case_data):
                evidence_list.append({
                    "name": rule.name,
                    "category": rule.category,
                    "description": rule.description,
                    "evidence_type": rule.evidence_type,
                    "copies": rule.copies,
                    "pages": rule.pages,
                })
        
        return evidence_list
    
    def to_text(self, evidence_list: List[Dict[str, Any]]) -> str:
        """
        生成文本格式证据清单
        
        Args:
            evidence_list: 证据列表
            
        Returns:
            格式化的文本字符串
        """
        if not evidence_list:
            return "暂无证据清单"
        
        lines = [
            "=" * 60,
            "证  据  清  单",
            "=" * 60,
            "",
        ]
        
        # 必备证据
        required = [e for e in evidence_list if e["category"] == "必备"]
        if required:
            lines.append("【必备证据】")
            lines.append("-" * 40)
            for i, evidence in enumerate(required, 1):
                lines.append(f"{i}. {evidence['name']}")
                lines.append(f"   证据形式：{evidence['evidence_type']}")
                lines.append(f"   份数/页数：{evidence['copies']}份/{evidence['pages']}页")
                lines.append(f"   证明目的：{evidence['description']}")
                lines.append("")
        
        # 建议补充证据
        optional = [e for e in evidence_list if e["category"] == "建议补充"]
        if optional:
            lines.append("【建议补充证据】")
            lines.append("-" * 40)
            for i, evidence in enumerate(optional, 1):
                lines.append(f"{i}. {evidence['name']}")
                lines.append(f"   证据形式：{evidence['evidence_type']}")
                lines.append(f"   份数/页数：{evidence['copies']}份/{evidence['pages']}页")
                lines.append(f"   证明目的：{evidence['description']}")
                lines.append("")
        
        lines.append("=" * 60)
        lines.append(f"共计 {len(evidence_list)} 项证据")
        lines.append(f"  必备证据：{len(required)} 项")
        lines.append(f"  建议补充：{len(optional)} 项")
        lines.append("=" * 60)
        
        return "\n".join(lines)
    
    def to_docx(self, evidence_list: List[Dict[str, Any]], 
                output_path: str, 
                case_info: Optional[Dict[str, str]] = None) -> str:
        """
        生成Word格式证据清单
        
        6列表格：序号/证据名称/份数/页数/证据形式/证明目的
        
        Args:
            evidence_list: 证据列表
            output_path: 输出文件路径
            case_info: 案件信息（可选）
                示例: {"案号": "(2024)冀01民初123号", "当事人": "张三诉李四"}
            
        Returns:
            输出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        doc = Document()
        
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("证 据 清 单")
        run.bold = True
        run.font.size = Pt(16)
        
        # 案件信息
        if case_info:
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for key, value in case_info.items():
                run = info_para.add_run(f"{key}：{value}    ")
                run.font.size = Pt(12)
            doc.add_paragraph()  # 空行
        
        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格样式
        table.style = 'Table Grid'
        
        # 表头
        headers = ["序号", "证据名称", "份数", "页数", "证据形式", "证明目的"]
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 设置表头格式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        
        # 填充数据行
        for idx, evidence in enumerate(evidence_list, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = evidence["name"]
            row.cells[2].text = str(evidence["copies"])
            row.cells[3].text = str(evidence["pages"])
            row.cells[4].text = evidence["evidence_type"]
            row.cells[5].text = evidence["description"]
            
            # 设置数据行格式
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # 设置列宽
        column_widths = [Cm(1.5), Cm(4), Cm(1.5), Cm(1.5), Cm(2.5), Cm(6)]
        for row in table.rows:
            for i, width in enumerate(column_widths):
                row.cells[i].width = width
        
        # 页脚说明
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = footer.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(9)
        run.font.color.rgb = None  # 默认颜色
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def create_common_rules() -> 'EvidenceEngine':
        """
        创建常用证据规则（交通事故案件）
        
        Returns:
            预配置的EvidenceEngine实例
        """
        engine = EvidenceEngine()
        
        # 身份证明类
        identity_rules = [
            EvidenceRule(
                name="原告身份证复印件",
                category="必备",
                condition=lambda d: True,
                description="证明原告身份信息",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="原告户口本复印件",
                category="必备",
                condition=lambda d: True,
                description="证明原告户籍性质（城镇/农村）",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="被告身份信息",
                category="必备",
                condition=lambda d: True,
                description="证明被告身份信息",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="被告驾驶证复印件",
                category="建议补充",
                condition=lambda d: d.get("has_driver_license", True),
                description="证明被告驾驶资格",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="被告行驶证复印件",
                category="建议补充",
                condition=lambda d: d.get("has_vehicle_registration", True),
                description="证明肇事车辆登记信息",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(identity_rules)
        
        # 事故证明类
        accident_rules = [
            EvidenceRule(
                name="交通事故认定书",
                category="必备",
                condition=lambda d: True,
                description="证明事故发生经过及责任划分",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="事故现场照片",
                category="建议补充",
                condition=lambda d: True,
                description="证明事故现场情况",
                evidence_type="视听资料"
            ),
            EvidenceRule(
                name="事故现场监控视频",
                category="建议补充",
                condition=lambda d: d.get("has_surveillance", False),
                description="证明事故发生经过",
                evidence_type="视听资料"
            ),
        ]
        engine.register_batch(accident_rules)
        
        # 医疗类
        medical_rules = [
            EvidenceRule(
                name="门诊病历",
                category="必备",
                condition=lambda d: d.get("has_medical_treatment", True),
                description="证明原告就医情况",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="住院病历",
                category="必备",
                condition=lambda d: d.get("hospitalized", False),
                description="证明原告住院治疗情况",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="医疗费发票",
                category="必备",
                condition=lambda d: d.get("has_medical_expense", True),
                description="证明原告医疗费用支出",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="用药清单",
                category="建议补充",
                condition=lambda d: d.get("hospitalized", False),
                description="证明用药情况及费用合理性",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="出院小结",
                category="建议补充",
                condition=lambda d: d.get("hospitalized", False),
                description="证明出院时身体状况及后续治疗建议",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(medical_rules)
        
        # 伤残鉴定类
        disability_rules = [
            EvidenceRule(
                name="司法鉴定意见书",
                category="必备",
                condition=lambda d: d.get("disability_grade", 0) > 0,
                description="证明原告伤残等级",
                evidence_type="鉴定意见"
            ),
            EvidenceRule(
                name="鉴定费发票",
                category="建议补充",
                condition=lambda d: d.get("disability_grade", 0) > 0,
                description="证明鉴定费用支出",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="后续治疗费鉴定意见",
                category="建议补充",
                condition=lambda d: d.get("has_followup_treatment", False),
                description="证明后续治疗费用",
                evidence_type="鉴定意见"
            ),
            EvidenceRule(
                name="护理依赖程度鉴定",
                category="建议补充",
                condition=lambda d: d.get("disability_grade", 0) >= 4,
                description="证明护理依赖程度",
                evidence_type="鉴定意见"
            ),
        ]
        engine.register_batch(disability_rules)
        
        # 误工/收入类
        income_rules = [
            EvidenceRule(
                name="劳动合同",
                category="建议补充",
                condition=lambda d: d.get("has_employment", False),
                description="证明原告工作情况",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="工资流水",
                category="建议补充",
                condition=lambda d: d.get("has_employment", False),
                description="证明原告收入情况",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="误工证明",
                category="建议补充",
                condition=lambda d: d.get("has_lost_income", False),
                description="证明原告因事故误工",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="纳税证明",
                category="建议补充",
                condition=lambda d: d.get("income_amount", 0) > 10000,
                description="证明原告收入情况",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(income_rules)
        
        # 护理/交通类
        nursing_rules = [
            EvidenceRule(
                name="护理费发票/收据",
                category="建议补充",
                condition=lambda d: d.get("nursing_days", 0) > 0,
                description="证明护理费用支出",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="护理人员身份证明",
                category="建议补充",
                condition=lambda d: d.get("nursing_days", 0) > 0,
                description="证明护理人员身份",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="交通费票据",
                category="建议补充",
                condition=lambda d: d.get("transport_days", 0) > 0,
                description="证明交通费用支出",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="住宿费发票",
                category="建议补充",
                condition=lambda d: d.get("accommodation_days", 0) > 0,
                description="证明住宿费用支出",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(nursing_rules)
        
        # 财产损失类
        property_rules = [
            EvidenceRule(
                name="车辆定损单",
                category="建议补充",
                condition=lambda d: d.get("vehicle_damage", 0) > 0,
                description="证明车辆损失金额",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="车辆维修发票",
                category="建议补充",
                condition=lambda d: d.get("vehicle_damage", 0) > 0,
                description="证明车辆维修费用",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="物品损失清单及照片",
                category="建议补充",
                condition=lambda d: d.get("property_damage", 0) > 0,
                description="证明其他财产损失",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(property_rules)
        
        # 被扶养人生活费类
        dependent_rules = [
            EvidenceRule(
                name="被扶养人身份证明",
                category="必备",
                condition=lambda d: d.get("has_dependents", False),
                description="证明被扶养人身份",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="亲属关系证明",
                category="必备",
                condition=lambda d: d.get("has_dependents", False),
                description="证明被扶养人与受害人关系",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="被扶养人无劳动能力证明",
                category="建议补充",
                condition=lambda d: d.get("dependent_type") in ["老人", "残疾人"],
                description="证明被扶养人无劳动能力",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(dependent_rules)
        
        # 保险类
        insurance_rules = [
            EvidenceRule(
                name="交强险保单",
                category="必备",
                condition=lambda d: d.get("has_compulsory_insurance", True),
                description="证明肇事车辆投保交强险情况",
                evidence_type="书证"
            ),
            EvidenceRule(
                name="商业三者险保单",
                category="建议补充",
                condition=lambda d: d.get("has_commercial_insurance", False),
                description="证明肇事车辆投保商业险情况",
                evidence_type="书证"
            ),
        ]
        engine.register_batch(insurance_rules)
        
        return engine
